# -*- coding: utf-8 -*-

# -----------------------------------------------------------------------------
# Nombre de archivo: flask_app.py
# Este es el archivo principal de la aplicación para PythonAnywhere.
# -----------------------------------------------------------------------------

import os
import io
import re
from datetime import datetime, timedelta, date

from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from flask_sqlalchemy import SQLAlchemy

# --- Configuración de la Aplicación Flask ---
app = Flask(__name__)
# La SECRET_KEY es importante para las sesiones de usuario y debe ser secreta en un proyecto real.
app.config['SECRET_KEY'] = 'una-clave-secreta-muy-dificil-de-adivinar-para-hac2025'

# --- Configuración de la Base de Datos (MySQL para PythonAnywhere) ---
# El código ahora toma las credenciales de las variables de entorno de PythonAnywhere.
# Esto es más seguro y evita errores de tipeo.
username = os.environ.get('PA_USER')
password = os.environ.get('PA_DB_PASSWORD')
hostname = os.environ.get('PA_DB_HOSTNAME')
databasename = os.environ.get('PA_DB_NAME')

# Si el código se está ejecutando en PythonAnywhere, usa la base de datos de ahí.
if username and password and hostname and databasename:
    app.config['SQLALCHEMY_DATABASE_URI'] = f"mysql+mysqlconnector://{username}:{password}@{hostname}/{databasename}"
    app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {'pool_recycle': 280}  # Requerido por PythonAnywhere
else:
    # Si no, usa una base de datos local para pruebas.
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///local_test.db'

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)


# --- Modelos de la Base de Datos (SQLAlchemy ORM) ---
class TimeEntry(db.Model):
    __tablename__ = 'time_entry'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=False)
    clock_in = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    clock_out = db.Column(db.DateTime, nullable=True)


class BalanceAdjustment(db.Model):
    __tablename__ = 'balance_adjustment'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=False)
    adjustment_seconds = db.Column(db.Integer, nullable=False)
    reason = db.Column(db.Text, nullable=False)
    admin_user_id = db.Column(db.Integer, nullable=False)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)


class Inventory(db.Model):
    __tablename__ = 'inventory'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150), nullable=False)
    brand = db.Column(db.String(100))
    color = db.Column(db.String(100))
    quantity = db.Column(db.String(50))
    status = db.Column(db.String(50))
    location = db.Column(db.String(150))


# --- Configuración de Flask-Login ---
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = "Por favor, inicia sesión para acceder a esta página."
login_manager.login_message_category = "info"


class User(UserMixin):
    def __init__(self, id, username, password_hash, daily_hours):
        self.id = id
        self.username = username
        self.password = password_hash
        self.daily_hours = daily_hours


USERS = {
    '1': {'username': 'Admin', 'password': generate_password_hash('Rindy'), 'daily_hours': 0},
    '2': {'username': 'Esteban', 'password': generate_password_hash('HAC127@'), 'daily_hours': 9.6},
    '3': {'username': 'Hernan', 'password': generate_password_hash('Obi123'), 'daily_hours': 9.6},
    '4': {'username': 'Jervaice', 'password': generate_password_hash('Tortas'), 'daily_hours': 8.0}
}


@login_manager.user_loader
def load_user(user_id):
    user_data = USERS.get(str(user_id))
    if user_data:
        return User(user_id, user_data['username'], user_data['password'], user_data['daily_hours'])
    return None


# --- Lógica de Negocio y Funciones Auxiliares ---
def get_user_by_username(username):
    for user_id, user_data in USERS.items():
        if user_data['username'] == username:
            return User(user_id, user_data['username'], user_data['password'], user_data['daily_hours'])
    return None


def seconds_to_str(seconds):
    if seconds is None: return ""
    sign = "-" if seconds < 0 else ""
    seconds = abs(int(seconds))
    hours, remainder = divmod(seconds, 3600)
    minutes, _ = divmod(remainder, 60)
    return f"{sign}{int(hours):02d}h {int(minutes):02d}m"


def parse_balance_string(balance_str):
    balance_str = balance_str.strip()
    match = re.match(r'([+-])?\s*(\d+):(\d+)', balance_str)
    if not match:
        raise ValueError("Formato de balance inválido. Use +HH:MM o -HH:MM.")
    sign, hours, minutes = match.groups()
    total_seconds = int(hours) * 3600 + int(minutes) * 60
    return -total_seconds if sign == '-' else total_seconds


def calculate_balance(user_id):
    uid = int(user_id)
    user = load_user(str(uid))
    if not user or user.daily_hours == 0:
        return "N/A", 0

    entries = TimeEntry.query.filter_by(user_id=uid).filter(TimeEntry.clock_out.isnot(None)).all()
    total_worked_seconds = sum((e.clock_out - e.clock_in).total_seconds() for e in entries)

    work_days = {e.clock_in.date() for e in entries if e.clock_in.weekday() < 5}
    required_seconds = len(work_days) * user.daily_hours * 3600

    adjustments = db.session.query(db.func.sum(BalanceAdjustment.adjustment_seconds)).filter_by(
        user_id=uid).scalar() or 0

    balance_seconds = (total_worked_seconds - required_seconds) + adjustments

    return seconds_to_str(balance_seconds), balance_seconds


# --- Rutas de la Aplicación (Controladores) ---
@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('login.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = get_user_by_username(username)
        if user and check_password_hash(user.password, password):
            login_user(user)
            flash(f'Bienvenido de nuevo, {user.username}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Usuario o contraseña incorrectos.', 'danger')
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Has cerrado sesión correctamente.', 'info')
    return redirect(url_for('login'))


@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')


@app.route('/time_tracking')
@login_required
def time_tracking():
    if current_user.username == 'Admin':
        return redirect(url_for('admin_view'))

    uid = int(current_user.id)
    entries = TimeEntry.query.filter_by(user_id=uid).order_by(TimeEntry.clock_in.desc()).all()
    active_entry = TimeEntry.query.filter_by(user_id=uid, clock_out=None).first()
    balance_str, balance_seconds = calculate_balance(uid)

    return render_template('time_tracking.html', entries=entries, active_entry=active_entry, balance_str=balance_str,
                           balance_seconds=balance_seconds)


@app.route('/clock_in', methods=['POST'])
@login_required
def clock_in():
    new_entry = TimeEntry(user_id=int(current_user.id), clock_in=datetime.utcnow())
    db.session.add(new_entry)
    db.session.commit()
    flash('Entrada marcada correctamente.', 'success')
    return redirect(url_for('time_tracking'))


@app.route('/clock_out', methods=['POST'])
@login_required
def clock_out():
    entry = TimeEntry.query.filter_by(user_id=int(current_user.id), clock_out=None).first()
    if entry:
        entry.clock_out = datetime.utcnow()
        db.session.commit()
        flash('Salida marcada correctamente.', 'success')
    return redirect(url_for('time_tracking'))


@app.route('/forgot_clock', methods=['GET'])
@login_required
def forgot_clock():
    if current_user.username not in ['Esteban', 'Hernan']:
        flash('No tienes permiso para acceder a esta función.', 'danger')
        return redirect(url_for('time_tracking'))
    return render_template('forgot_clock.html')


@app.route('/save_forgotten_clock', methods=['POST'])
@login_required
def save_forgotten_clock():
    if current_user.username not in ['Esteban', 'Hernan']:
        return redirect(url_for('time_tracking'))
    entry_date = request.form['date']
    clock_in_time = request.form['clock_in_time']
    clock_out_time = request.form['clock_out_time']
    try:
        clock_in_dt = datetime.strptime(f"{entry_date} {clock_in_time}", '%Y-%m-%d %H:%M')
        clock_out_dt = datetime.strptime(f"{entry_date} {clock_out_time}", '%Y-%m-%d %H:%M')
        if clock_out_dt <= clock_in_dt:
            flash('La hora de salida debe ser posterior a la de entrada.', 'danger')
            return redirect(url_for('forgot_clock'))

        new_entry = TimeEntry(user_id=int(current_user.id), clock_in=clock_in_dt, clock_out=clock_out_dt)
        db.session.add(new_entry)
        db.session.commit()
        flash('Marcaje manual añadido correctamente.', 'success')
    except ValueError:
        flash('Formato de fecha u hora inválido.', 'danger')
        return redirect(url_for('forgot_clock'))
    return redirect(url_for('time_tracking'))


@app.route('/admin')
@login_required
def admin_view():
    if current_user.username != 'Admin':
        flash('Acceso no autorizado.', 'danger')
        return redirect(url_for('dashboard'))

    users_info = {}
    for user_id_str, user_data in USERS.items():
        if user_data['username'] == 'Admin': continue
        user_id = int(user_id_str)
        balance_str, balance_seconds = calculate_balance(user_id)
        entries = TimeEntry.query.filter_by(user_id=user_id).order_by(TimeEntry.clock_in.desc()).all()
        adjustments = BalanceAdjustment.query.filter_by(user_id=user_id).order_by(
            BalanceAdjustment.created_at.desc()).all()

        for adj in adjustments:
            adj.adjustment_str = seconds_to_str(adj.adjustment_seconds)
            admin_user = load_user(adj.admin_user_id)
            adj.admin_username = admin_user.username if admin_user else "Desconocido"

        users_info[user_id] = {'username': user_data['username'], 'balance_str': balance_str,
                               'balance_seconds': balance_seconds, 'entries': entries, 'adjustments': adjustments}

    return render_template('admin_view.html', users_info=users_info)


@app.route('/admin/edit/<int:entry_id>', methods=['POST'])
@login_required
def admin_edit_entry(entry_id):
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    entry = TimeEntry.query.get(entry_id)
    if entry:
        entry.clock_in = datetime.strptime(request.form.get('clock_in'), '%Y-%m-%dT%H:%M')
        clock_out_str = request.form.get('clock_out')
        entry.clock_out = datetime.strptime(clock_out_str, '%Y-%m-%dT%H:%M') if clock_out_str else None
        db.session.commit()
        flash('Marcaje actualizado.', 'success')
    return redirect(url_for('admin_view'))


@app.route('/admin/delete/<int:entry_id>')
@login_required
def admin_delete_entry(entry_id):
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    entry = TimeEntry.query.get(entry_id)
    if entry:
        db.session.delete(entry)
        db.session.commit()
        flash('Marcaje eliminado.', 'success')
    return redirect(url_for('admin_view'))


@app.route('/admin/set_balance/<int:user_id>', methods=['POST'])
@login_required
def admin_set_balance(user_id):
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    new_total_balance_str = request.form.get('new_total_balance')
    reason = request.form.get('reason')
    if not reason:
        flash('Se requiere un motivo para el ajuste de balance.', 'danger')
        return redirect(url_for('admin_view'))
    try:
        new_total_seconds = parse_balance_string(new_total_balance_str)
    except ValueError as e:
        flash(str(e), 'danger')
        return redirect(url_for('admin_view'))

    _, current_balance_seconds = calculate_balance(user_id)
    adjustment_needed = new_total_seconds - current_balance_seconds

    new_adj = BalanceAdjustment(user_id=user_id, adjustment_seconds=adjustment_needed, reason=reason,
                                admin_user_id=int(current_user.id), created_at=datetime.utcnow())
    db.session.add(new_adj)
    db.session.commit()
    flash(f"Balance para el usuario actualizado correctamente.", 'success')
    return redirect(url_for('admin_view'))


@app.route('/admin/letter_generator')
@login_required
def letter_generator():
    if current_user.username != 'Admin':
        flash('Acceso no autorizado.', 'danger')
        return redirect(url_for('dashboard'))
    today_date = date.today().strftime('%Y-%m-%d')
    current_year = date.today().year
    return render_template('letter_generator.html', today_date=today_date, current_year=current_year)


@app.route('/admin/generate_doc', methods=['POST'])
@login_required
def generate_doc():
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    data = {k: request.form.get(k) for k in
            ['doc_date', 'project_name', 'client_name', 'year', 'contact_person', 'contact_position', 'contact_email',
             'project_type']}
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    try:
        date_obj = datetime.strptime(data['doc_date'], '%Y-%m-%d')
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre",
                 "noviembre", "diciembre"]
        formatted_date = f"{date_obj.day} de {meses[date_obj.month - 1]} de {date_obj.year}"
    except (ValueError, TypeError):
        formatted_date = data['doc_date']
    p_date = document.add_paragraph(f'Fecha: {formatted_date}')
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()
    document.add_paragraph('A quien interese:', style='Normal').add_run().bold = True
    document.add_paragraph(
        'Por medio de la presente, queremos expresar nuestra satisfacción con los trabajos realizados por HidroAmbiente Consultores.')
    document.add_paragraph(
        'Su profesionalismo y compromiso han sido fundamentales para el éxito de los proyectos desarrollados, cumpliendo con nuestras expectativas y necesidades.')
    document.add_paragraph('A continuación, se presenta un cuadro detallado con la información correspondiente:')
    document.add_paragraph()
    table_data = [("Tipo de Proyecto", data['project_type']), ("Nombre del Proyecto", data['project_name']),
                  ("Nombre del Cliente", data['client_name']), ("Año", data['year']),
                  ("Contacto", data['contact_person']), ("Puesto del Contacto", data['contact_position']),
                  ("Correo Electrónico", data['contact_email'])]
    table = document.add_table(rows=len(table_data), cols=2)
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(table_data):
        table.cell(i, 0).text = key
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value
    document.add_paragraph()
    document.add_paragraph(
        'Finalmente, agradecemos nuevamente su excelente trabajo y quedamos en espera de futuras colaboraciones que sigan fortaleciendo esta relación profesional.')
    document.add_paragraph()
    document.add_paragraph('Atentamente,')
    document.add_paragraph('\n\n')
    document.add_paragraph('_________________________')
    document.add_paragraph(data['contact_person']).add_run().bold = True
    document.add_paragraph(data['contact_position'])
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name=f"Carta_Satisfaccion_{data['project_name']}.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# --- Rutas para Inventario ---
@app.route('/inventory')
@login_required
def inventory_view():
    if current_user.username not in ['Admin', 'Jervaice']:
        flash('Acceso no autorizado.', 'danger')
        return redirect(url_for('dashboard'))
    items = Inventory.query.order_by(Inventory.name).all()
    return render_template('inventory.html', items=items)


@app.route('/inventory/add', methods=['POST'])
@login_required
def add_inventory_item():
    if current_user.username not in ['Admin', 'Jervaice']: return redirect(url_for('dashboard'))
    new_item = Inventory(
        name=request.form.get('name'), brand=request.form.get('brand'), color=request.form.get('color'),
        quantity=request.form.get('quantity'), status=request.form.get('status'), location=request.form.get('location')
    )
    db.session.add(new_item)
    db.session.commit()
    flash('Activo añadido correctamente.', 'success')
    return redirect(url_for('inventory_view'))


@app.route('/inventory/edit/<int:item_id>', methods=['POST'])
@login_required
def edit_inventory_item(item_id):
    if current_user.username not in ['Admin', 'Jervaice']: return redirect(url_for('dashboard'))
    item = Inventory.query.get(item_id)
    if item:
        item.name = request.form.get('name')
        item.brand = request.form.get('brand')
        item.color = request.form.get('color')
        item.quantity = request.form.get('quantity')
        item.status = request.form.get('status')
        item.location = request.form.get('location')
        db.session.commit()
        flash('Activo actualizado correctamente.', 'success')
    return redirect(url_for('inventory_view'))


@app.route('/inventory/delete/<int:item_id>')
@login_required
def delete_inventory_item(item_id):
    if current_user.username not in ['Admin', 'Jervaice']: return redirect(url_for('dashboard'))
    item = Inventory.query.get(item_id)
    if item:
        db.session.delete(item)
        db.session.commit()
        flash('Activo eliminado correctamente.', 'success')
    return redirect(url_for('inventory_view'))


@app.route('/inventory/export', methods=['POST'])
@login_required
def export_inventory():
    if current_user.username not in ['Admin', 'Jervaice']: return redirect(url_for('dashboard'))

    export_type = request.form.get('export_type')

    if export_type == 'selected':
        selected_ids_str = request.form.get('selected_ids', '').split(',')
        if not selected_ids_str or selected_ids_str == ['']:
            flash('No se seleccionaron activos para exportar.', 'warning')
            return redirect(url_for('inventory_view'))
        selected_ids = [int(i) for i in selected_ids_str if i]
        items = Inventory.query.filter(Inventory.id.in_(selected_ids)).order_by(Inventory.name).all()
    else:
        items = Inventory.query.order_by(Inventory.name).all()

    wb = Workbook()
    ws = wb.active
    ws.title = "Inventario de Activos"
    headers = ["Activo", "Marca", "Color", "Cantidad", "Estado", "Ubicación"]
    ws.append(headers)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
    for item in items:
        ws.append([item.name, item.brand, item.color, item.quantity, item.status, item.location])

    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream, as_attachment=True,
        download_name=f"Inventario_HAC2025_{date.today().isoformat()}.xlsx",
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


# --- Ruta de Inicialización de la Base de Datos ---
@app.route('/init-db')
def init_db_command():
    try:
        with app.app_context():
            db.create_all()
            if Inventory.query.count() == 0:
                initial_data = get_initial_inventory_data()
                for item_data in initial_data:
                    db.session.add(Inventory(**item_data))
                db.session.commit()
                return "Base de datos inicializada y poblada."
            else:
                return "La base de datos ya existe y contiene datos."
    except Exception as e:
        return f"Ocurrió un error al inicializar la base de datos: {e}"


def get_initial_inventory_data():
    return [
        {'name': 'Botas de Hule', 'brand': 'Varios', 'color': 'Negro', 'quantity': '7', 'status': 'Bueno',
         'location': 'Estante Cochera Izquierda'},
        # ... (todos los demás datos del inventario van aquí)
    ]
