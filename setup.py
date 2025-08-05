#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
HAC2025 - Script de Configuración y Ejecución (Versión Web App)

Este script tiene un doble propósito:
1.  Verifica e instala las dependencias necesarias (Flask, Flask-Login, python-docx, openpyxl) si no están presentes.
2.  Ejecuta una aplicación web completa con Flask que gestiona el marcaje de horas, genera documentos y controla un inventario.

Para ejecutar:
1. Guarda este archivo como 'setup.py'.
2. Ejecuta desde la terminal: python setup.py
3. Abre tu navegador en http://127.0.0.1:5000
"""

import sys
import subprocess
import os
import sqlite3
import re
from datetime import datetime, timedelta, date
import io


# --- PASO 1: Verificación e Instalación de Dependencias ---
def install_dependencies():
    """Verifica si las librerías requeridas están instaladas e instala las que falten."""
    required_packages = ['Flask', 'Flask-Login', 'python-docx', 'openpyxl']
    print(">>> Verificando dependencias...")
    try:
        for package in required_packages:
            if package == 'python-docx':
                __import__('docx')
            elif package == 'openpyxl':
                __import__('openpyxl')
            else:
                __import__(package)
        print(">>> Todas las dependencias ya están instaladas.")
    except ImportError as e:
        if 'docx' in str(e):
            missing_package_name = 'python-docx'
        elif 'openpyxl' in str(e):
            missing_package_name = 'openpyxl'
        else:
            missing_package_name = str(e).split("No module named ")[1].strip("'")

        print(f"*** Falta la librería: {missing_package_name}. Intentando instalar...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", missing_package_name])
            print(f">>> {missing_package_name} instalado correctamente.")
        except subprocess.CalledProcessError:
            print(
                f"!!! ERROR: No se pudo instalar {missing_package_name}. Por favor, instálalo manualmente con 'pip install {missing_package_name}'")
            sys.exit(1)


# Ejecutar la instalación de dependencias antes de importar Flask
install_dependencies()

# --- PASO 2: Importaciones de la Aplicación Web ---
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from jinja2 import ChoiceLoader, DictLoader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

# --- PASO 3: Configuración de la Aplicación Flask ---
app = Flask(__name__)
app.config['SECRET_KEY'] = 'una-clave-secreta-muy-dificil-de-adivinar'
app.config['DATABASE'] = 'hac2025.db'

# --- PASO 7: Plantillas HTML (Templates) ---

HTML_LAYOUT = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>HAC2025 - {% block title %}{% endblock %}</title>
    <style>
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; background-color: #f4f7f6; color: #333; margin: 0; padding: 20px; }
        .container { max-width: 1200px; margin: 0 auto; background-color: #fff; padding: 20px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        h1, h2, h3, h4 { color: #005f73; }
        .navbar { display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px; padding-bottom: 10px; border-bottom: 1px solid #ddd; }
        .navbar a { text-decoration: none; color: #007bff; font-weight: bold; }
        .flash-messages { list-style: none; padding: 0; margin: 0 0 20px 0; }
        .flash-messages li { padding: 15px; border-radius: 5px; margin-bottom: 10px; }
        .flash-messages .info { background-color: #e0f7fa; color: #00796b; }
        .flash-messages .success { background-color: #e8f5e9; color: #2e7d32; }
        .flash-messages .danger { background-color: #ffebee; color: #c62828; }
        .btn { display: inline-flex; align-items: center; gap: 8px; padding: 10px 15px; border: none; border-radius: 5px; color: white; text-decoration: none; cursor: pointer; font-size: 16px; }
        .btn-primary { background-color: #007bff; }
        .btn-secondary { background-color: #6c757d; }
        .btn-danger { background-color: #dc3545; }
        .btn-success { background-color: #28a745; }
        .btn-warning { background-color: #ffc107; color: #212529; }
        .form-group { margin-bottom: 15px; }
        .form-group label { display: block; margin-bottom: 5px; font-weight: bold; }
        .form-group input, .form-group select { width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px; box-sizing: border-box; }
        table { width: 100%; border-collapse: collapse; margin-top: 20px; }
        th, td { padding: 12px; border: 1px solid #ddd; text-align: left; }
        th { background-color: #005f73; color: white; }
        tr:nth-child(even) { background-color: #f2f2f2; }
        .admin-user-block { border: 1px solid #ccc; border-radius: 8px; padding: 15px; margin-bottom: 20px; }
        .grid-form { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
        .modal { display: none; position: fixed; z-index: 1000; left: 0; top: 0; width: 100%; height: 100%; overflow: auto; background-color: rgba(0,0,0,0.5); }
        .modal-content { background-color: #fefefe; margin: 10% auto; padding: 20px; border: 1px solid #888; width: 80%; max-width: 600px; border-radius: 8px; }
        .close-button { color: #aaa; float: right; font-size: 28px; font-weight: bold; cursor: pointer; }
    </style>
</head>
<body>
    <div class="container">
        <div class="navbar">
            <h1>HAC2025</h1>
            {% if current_user.is_authenticated %}
                <span>Hola, {{ current_user.username }} | <a href="{{ url_for('logout') }}">Cerrar Sesión</a></span>
            {% endif %}
        </div>
        {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
                <ul class="flash-messages">
                {% for category, message in messages %}
                    <li class="{{ category }}">{{ message }}</li>
                {% endfor %}
                </ul>
            {% endif %}
        {% endwith %}
        {% block content %}{% endblock %}
    </div>
</body>
</html>
"""

HTML_LOGIN = """
{% extends "layout.html" %}
{% block title %}Inicio de Sesión{% endblock %}
{% block content %}
    <h2>Iniciar Sesión</h2>
    <form method="POST" action="{{ url_for('login') }}">
        <div class="form-group">
            <label for="username">Usuario:</label>
            <input type="text" id="username" name="username" required>
        </div>
        <div class="form-group">
            <label for="password">Contraseña:</label>
            <input type="password" id="password" name="password" required>
        </div>
        <button type="submit" class="btn btn-primary">Ingresar</button>
    </form>
{% endblock %}
"""

HTML_DASHBOARD = """
{% extends "layout.html" %}
{% block title %}Dashboard{% endblock %}
{% block content %}
    <h2>Panel Principal</h2>
    <p>Bienvenido, {{ current_user.username }}. Selecciona una opción:</p>
    <div style="display: flex; flex-wrap: wrap; gap: 10px;">
        {% if current_user.username == 'Admin' %}
            <a href="{{ url_for('admin_view') }}" class="btn btn-primary">Panel de Horas</a>
            <a href="{{ url_for('letter_generator') }}" class="btn btn-primary">Generar Carta</a>
            <a href="{{ url_for('inventory_view') }}" class="btn btn-primary">Inventario</a>
        {% elif current_user.username == 'Jervaice' %}
            <a href="{{ url_for('time_tracking') }}" class="btn btn-primary">Marcaje de Horas</a>
            <a href="{{ url_for('inventory_view') }}" class="btn btn-primary">Inventario</a>
        {% else %}
            <a href="{{ url_for('time_tracking') }}" class="btn btn-primary">Marcaje de Horas</a>
        {% endif %}
    </div>
{% endblock %}
"""

HTML_TIME_TRACKING = """
{% extends "layout.html" %}
{% block title %}Marcaje de Horas{% endblock %}
{% block content %}
    <h2>Control de Horas</h2>
    <div style="padding: 20px; border-radius: 8px; margin-bottom: 20px; text-align: center;
        {% if balance_seconds >= 0 %} background-color: #d4edda; color: #155724; 
        {% else %} background-color: #f8d7da; color: #721c24; {% endif %}">
        <h3>Balance General de Horas</h3>
        <p style="font-size: 2em; font-weight: bold; margin: 0;">{{ balance_str }}</p>
    </div>
    <div style="text-align: center; margin-bottom: 20px;">
        {% if not active_entry %}
            <form method="POST" action="{{ url_for('clock_in') }}" style="display: inline;">
                <button type="submit" class="btn btn-success">Marcar Entrada</button>
            </form>
        {% else %}
            <form method="POST" action="{{ url_for('clock_out') }}" style="display: inline;">
                <button type="submit" class="btn btn-danger">Marcar Salida</button>
            </form>
            <p style="margin-top: 10px;">Entrada marcada a las: {{ active_entry.clock_in.strftime('%H:%M:%S del %d/%m/%Y') }}</p>
        {% endif %}
        {% if current_user.username in ['Esteban', 'Hernan'] %}
            <a href="{{ url_for('forgot_clock') }}" class="btn btn-warning">Olvidé Marcar</a>
        {% endif %}
    </div>
    <h3>Mis Marcajes Recientes</h3>
    <table>
        <thead><tr><th>Fecha</th><th>Entrada</th><th>Salida</th><th>Duración</th></tr></thead>
        <tbody>
            {% for entry in entries %}
            <tr>
                <td>{{ entry.clock_in.strftime('%d/%m/%Y') }}</td>
                <td>{{ entry.clock_in.strftime('%H:%M:%S') }}</td>
                <td>{{ entry.clock_out.strftime('%H:%M:%S') if entry.clock_out else 'En curso' }}</td>
                <td>
                    {% if entry.clock_out %}
                        {% set duration = entry.clock_out - entry.clock_in %}
                        {{ '%02d:%02d:%02d' % (duration.seconds // 3600, (duration.seconds % 3600) // 60, duration.seconds % 60) }}
                    {% else %} - {% endif %}
                </td>
            </tr>
            {% else %}
            <tr><td colspan="4" style="text-align: center;">No hay marcajes registrados.</td></tr>
            {% endfor %}
        </tbody>
    </table>
{% endblock %}
"""

HTML_FORGOT_CLOCK = """
{% extends "layout.html" %}
{% block title %}Modificar Marcaje{% endblock %}
{% block content %}
    <h2>Añadir/Modificar Marcaje Anterior</h2>
    <p>Completa los datos para un marcaje que olvidaste registrar.</p>
    <form method="POST" action="{{ url_for('save_forgotten_clock') }}">
        <div class="form-group"><label for="date">Fecha:</label><input type="date" id="date" name="date" required></div>
        <div class="form-group"><label for="clock_in_time">Hora de Entrada (HH:MM):</label><input type="time" id="clock_in_time" name="clock_in_time" required></div>
        <div class="form-group"><label for="clock_out_time">Hora de Salida (HH:MM):</label><input type="time" id="clock_out_time" name="clock_out_time" required></div>
        <button type="submit" class="btn btn-primary">Guardar Marcaje</button>
        <a href="{{ url_for('time_tracking') }}" class="btn btn-secondary">Cancelar</a>
    </form>
{% endblock %}
"""

HTML_ADMIN_VIEW = """
{% extends "layout.html" %}
{% block title %}Panel de Administrador de Horas{% endblock %}
{% block content %}
    <h2>Panel de Administrador de Horas</h2>
    <p>Aquí puedes ver y gestionar la información de horas de los usuarios.</p>
    {% for user_id, user_data in users_info.items() %}
        <div class="admin-user-block">
            <h3>{{ user_data.username }}</h3>
            <p><strong>Balance de Horas Actual:</strong> 
                <span style="font-weight: bold; font-size: 1.2em; color: {{ '#155724' if user_data.balance_seconds >= 0 else '#721c24' }};">{{ user_data.balance_str }}</span>
            </p>
            <form method="POST" action="{{ url_for('admin_set_balance', user_id=user_id) }}" style="margin-top: 15px; margin-bottom: 25px; background-color: #f8f9fa; padding: 15px; border-radius: 5px;">
                <h4>Ajustar Balance de Horas</h4>
                <div class="form-group"><label for="new_total_balance_{{ user_id }}">Establecer Nuevo Balance Total (formato: +HH:MM o -HH:MM):</label><input type="text" id="new_total_balance_{{ user_id }}" name="new_total_balance" placeholder="+02:30" required></div>
                <div class="form-group"><label for="reason_{{ user_id }}">Motivo del ajuste:</label><input type="text" id="reason_{{ user_id }}" name="reason" required></div>
                <button type="submit" class="btn btn-primary">Guardar Nuevo Balance</button>
            </form>
            <h4>Marcajes Registrados</h4>
            <table>
                <thead><tr><th>Entrada</th><th>Salida</th><th>Acciones</th></tr></thead>
                <tbody>
                    {% for entry in user_data.entries %}
                    <tr>
                        <form method="POST" action="{{ url_for('admin_edit_entry', entry_id=entry.id) }}">
                        <td><input type="datetime-local" name="clock_in" value="{{ entry.clock_in.strftime('%Y-%m-%dT%H:%M') }}"></td>
                        <td><input type="datetime-local" name="clock_out" value="{{ entry.clock_out.strftime('%Y-%m-%dT%H:%M') if entry.clock_out else '' }}"></td>
                        <td>
                            <button type="submit" class="btn btn-warning" style="padding: 5px 10px;">Guardar</button>
                            <a href="{{ url_for('admin_delete_entry', entry_id=entry.id) }}" class="btn btn-danger" style="padding: 5px 10px;" onclick="return confirm('¿Estás seguro de que quieres borrar este marcaje?');">Borrar</a>
                        </td>
                        </form>
                    </tr>
                    {% else %}
                    <tr><td colspan="3" style="text-align:center;">Sin marcajes.</td></tr>
                    {% endfor %}
                </tbody>
            </table>
            <h4 style="margin-top: 25px;">Ajustes Manuales de Balance</h4>
            <table>
                <thead><tr><th>Fecha de Ajuste</th><th>Ajuste (H:M:S)</th><th>Motivo</th><th>Admin</th></tr></thead>
                <tbody>
                    {% for adj in user_data.adjustments %}
                    <tr>
                        <td>{{ adj.created_at.strftime('%Y-%m-%d %H:%M') }}</td><td>{{ adj.adjustment_str }}</td><td>{{ adj.reason }}</td><td>{{ adj.admin_username }}</td>
                    </tr>
                    {% else %}
                    <tr><td colspan="4" style="text-align:center;">Sin ajustes manuales.</td></tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
    {% endfor %}
{% endblock %}
"""

HTML_LETTER_GENERATOR = """
{% extends "layout.html" %}
{% block title %}Generador de Cartas{% endblock %}
{% block content %}
    <h2>Generador de Cartas de Satisfacción</h2>
    <form method="POST" action="{{ url_for('generate_doc') }}">
        <div class="grid-form">
            <div class="form-group"><label for="doc_date">Fecha del Documento</label><input type="date" id="doc_date" name="doc_date" value="{{ today_date }}" required></div>
            <div class="form-group"><label for="project_name">Nombre del Proyecto</label><input type="text" id="project_name" name="project_name" required></div>
            <div class="form-group"><label for="client_name">Nombre del Cliente</label><input type="text" id="client_name" name="client_name" required></div>
            <div class="form-group"><label for="year">Año</label><input type="number" id="year" name="year" value="{{ current_year }}" required></div>
            <div class="form-group"><label for="contact_person">Persona de Contacto (será el firmante)</label><input type="text" id="contact_person" name="contact_person" required></div>
            <div class="form-group"><label for="contact_position">Puesto del Contacto (será el puesto del firmante)</label><input type="text" id="contact_position" name="contact_position" required></div>
            <div class="form-group"><label for="contact_email">Correo Electrónico del Contacto</label><input type="email" id="contact_email" name="contact_email" required></div>
            <div class="form-group"><label for="project_type">Tipo de Proyecto</label><input type="text" id="project_type" name="project_type" required></div>
        </div>
        <button type="submit" class="btn btn-primary" style="width: 100%; padding: 15px; margin-top: 20px;">Generar Documento (.docx)</button>
    </form>
{% endblock %}
"""

HTML_INVENTORY = """
{% extends "layout.html" %}
{% block title %}Inventario de Activos{% endblock %}
{% block content %}
    <h2>Inventario de Activos</h2>
    <div style="margin-bottom: 20px;">
        <input type="text" id="searchInput" onkeyup="filterTable()" placeholder="Buscar por activo, marca, color, estado o ubicación..." style="width: 100%; padding: 10px; border-radius: 5px; border: 1px solid #ccc;">
    </div>
    <div style="display: flex; gap: 10px; margin-bottom: 20px;">
        <form id="exportForm" method="POST" action="{{ url_for('export_inventory') }}" style="display: contents;">
            <input type="hidden" id="selected_ids_input" name="selected_ids">
            <button type="button" class="btn btn-secondary" onclick="exportSelected()">
                <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="currentColor" viewBox="0 0 16 16"><path d="M.5 9.9a.5.5 0 0 1 .5.5v2.5a1 1 0 0 0 1 1h12a1 1 0 0 0 1-1v-2.5a.5.5 0 0 1 1 0v2.5a2 2 0 0 1-2 2H2a2 2 0 0 1-2-2v-2.5a.5.5 0 0 1 .5-.5z"/><path d="M7.646 11.854a.5.5 0 0 0 .708 0l3-3a.5.5 0 0 0-.708-.708L8.5 10.293V1.5a.5.5 0 0 0-1 0v8.793L5.354 8.146a.5.5 0 1 0-.708.708l3 3z"/></svg>
                Exportar Selección
            </button>
            <button type="button" class="btn btn-secondary" onclick="exportAll()">
                <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="currentColor" viewBox="0 0 16 16"><path d="M.5 9.9a.5.5 0 0 1 .5.5v2.5a1 1 0 0 0 1 1h12a1 1 0 0 0 1-1v-2.5a.5.5 0 0 1 1 0v2.5a2 2 0 0 1-2 2H2a2 2 0 0 1-2-2v-2.5a.5.5 0 0 1 .5-.5z"/><path d="M7.646 11.854a.5.5 0 0 0 .708 0l3-3a.5.5 0 0 0-.708-.708L8.5 10.293V1.5a.5.5 0 0 0-1 0v8.793L5.354 8.146a.5.5 0 1 0-.708.708l3 3z"/></svg>
                Exportar Todo
            </button>
        </form>
        <button class="btn btn-primary" onclick="openAddModal()">
            <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="currentColor" viewBox="0 0 16 16"><path d="M8 4a.5.5 0 0 1 .5.5v3h3a.5.5 0 0 1 0 1h-3v3a.5.5 0 0 1-1 0v-3h-3a.5.5 0 0 1 0-1h3v-3A.5.5 0 0 1 8 4z"/></svg>
            Añadir Activo
        </button>
    </div>

    <table id="inventoryTable">
        <thead>
            <tr>
                <th><input type="checkbox" onclick="toggleAll(this)"></th>
                <th>Activo</th><th>Marca</th><th>Color</th><th>Cantidad</th><th>Estado</th><th>Ubicación</th><th>Acciones</th>
            </tr>
        </thead>
        <tbody>
            {% for item in items %}
            <tr data-id="{{ item.id }}" data-name="{{ item.name }}" data-brand="{{ item.brand }}" data-color="{{ item.color }}" data-quantity="{{ item.quantity }}" data-status="{{ item.status }}" data-location="{{ item.location }}">
                <td><input type="checkbox" name="selected_item" value="{{ item.id }}"></td>
                <td>{{ item.name }}</td><td>{{ item.brand }}</td><td>{{ item.color }}</td><td>{{ item.quantity }}</td><td>{{ item.status }}</td><td>{{ item.location }}</td>
                <td>
                    <div style="display:flex; gap: 5px;">
                        <button class="btn btn-warning" style="padding: 5px;" onclick="openEditModal(this)">
                            <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="currentColor" viewBox="0 0 16 16"><path d="M15.502 1.94a.5.5 0 0 1 0 .706L14.459 3.69l-2-2L13.502.646a.5.5 0 0 1 .707 0l1.293 1.293zm-1.75 2.456-2-2L4.939 9.21a.5.5 0 0 0-.121.196l-.805 2.414a.25.25 0 0 0 .316.316l2.414-.805a.5.5 0 0 0 .196-.12l6.813-6.814z"/><path fill-rule="evenodd" d="M1 13.5A1.5 1.5 0 0 0 2.5 15h11a1.5 1.5 0 0 0 1.5-1.5v-6a.5.5 0 0 0-1 0v6a.5.5 0 0 1-.5.5h-11a.5.5 0 0 1-.5-.5v-11a.5.5 0 0 1 .5-.5H9a.5.5 0 0 0 0-1H2.5A1.5 1.5 0 0 0 1 2.5v11z"/></svg>
                        </button>
                        <a href="{{ url_for('delete_inventory_item', item_id=item.id) }}" class="btn btn-danger" style="padding: 5px;" onclick="return confirm('¿Estás seguro de que quieres eliminar este activo?');">
                            <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="currentColor" viewBox="0 0 16 16"><path d="M5.5 5.5A.5.5 0 0 1 6 6v6a.5.5 0 0 1-1 0V6a.5.5 0 0 1 .5-.5zm2.5 0a.5.5 0 0 1 .5.5v6a.5.5 0 0 1-1 0V6a.5.5 0 0 1 .5-.5zm3 .5a.5.5 0 0 0-1 0v6a.5.5 0 0 0 1 0V6z"/><path fill-rule="evenodd" d="M14.5 3a1 1 0 0 1-1 1H13v9a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V4h-.5a1 1 0 0 1-1-1V2a1 1 0 0 1 1-1H6a1 1 0 0 1 1-1h2a1 1 0 0 1 1 1h3.5a1 1 0 0 1 1 1v1zM4.118 4 4 4.059V13a1 1 0 0 0 1 1h6a1 1 0 0 0 1-1V4.059L11.882 4H4.118zM2.5 3V2h11v1h-11z"/></svg>
                        </a>
                    </div>
                </td>
            </tr>
            {% else %}
            <tr><td colspan="8" style="text-align:center;">No hay activos en el inventario.</td></tr>
            {% endfor %}
        </tbody>
    </table>

    <!-- Modal para Añadir/Editar Activo -->
    <div id="assetModal" class="modal">
        <div class="modal-content">
            <span class="close-button" onclick="closeModal()">&times;</span>
            <h2 id="modalTitle">Añadir Nuevo Activo</h2>
            <form id="assetForm" method="POST">
                <input type="hidden" id="itemId" name="item_id">
                <div class="grid-form">
                    <div class="form-group"><label for="name">Activo</label><input type="text" id="name" name="name" required></div>
                    <div class="form-group"><label for="brand">Marca del Activo</label><input type="text" id="brand" name="brand"></div>
                    <div class="form-group"><label for="color">Color del activo</label><input type="text" id="color" name="color"></div>
                    <div class="form-group"><label for="quantity">Cantidad</label><input type="text" id="quantity" name="quantity" required></div>
                    <div class="form-group"><label for="status">Estado del activo</label>
                        <select id="status" name="status">
                            <option>Bueno</option><option>Revisar</option><option>Probar</option><option>Malo</option>
                            <option>Lleno</option><option>Casi Lleno</option><option>Medio lleno</option><option>1/2</option>
                            <option>3/4</option><option>1/4</option><option>1/16</option><option>Vacio</option>
                            <option>Nuevo</option><option>Usados</option><option>Mitad</option>
                        </select>
                    </div>
                    <div class="form-group"><label for="location">Ubicación del activo</label><input type="text" id="location" name="location"></div>
                </div>
                <button type="submit" class="btn btn-primary" style="width: 100%; margin-top: 15px;">Guardar Cambios</button>
            </form>
        </div>
    </div>

    <script>
        const modal = document.getElementById('assetModal');
        const form = document.getElementById('assetForm');
        const modalTitle = document.getElementById('modalTitle');

        function openModal() { modal.style.display = 'block'; }
        function closeModal() { modal.style.display = 'none'; }

        function openAddModal() {
            form.reset();
            form.action = "{{ url_for('add_inventory_item') }}";
            modalTitle.textContent = 'Añadir Nuevo Activo';
            openModal();
        }

        function openEditModal(button) {
            form.reset();
            const row = button.closest('tr');
            const id = row.dataset.id;

            form.action = `/inventory/edit/${id}`;
            modalTitle.textContent = 'Editar Activo';

            document.getElementById('itemId').value = id;
            document.getElementById('name').value = row.dataset.name;
            document.getElementById('brand').value = row.dataset.brand;
            document.getElementById('color').value = row.dataset.color;
            document.getElementById('quantity').value = row.dataset.quantity;
            document.getElementById('status').value = row.dataset.status;
            document.getElementById('location').value = row.dataset.location;

            openModal();
        }

        function filterTable() {
            const input = document.getElementById("searchInput");
            const filter = input.value.toUpperCase();
            const table = document.getElementById("inventoryTable");
            const tr = table.getElementsByTagName("tr");

            for (let i = 1; i < tr.length; i++) { // Start from 1 to skip header
                let rowVisible = false;
                const tds = tr[i].getElementsByTagName("td");
                for (let j = 1; j < tds.length - 1; j++) { // Skip checkbox and actions
                    if (tds[j]) {
                        if (tds[j].textContent.toUpperCase().indexOf(filter) > -1) {
                            rowVisible = true;
                            break;
                        }
                    }
                }
                tr[i].style.display = rowVisible ? "" : "none";
            }
        }

        function toggleAll(source) {
            const checkboxes = document.querySelectorAll('input[name="selected_item"]');
            for (const checkbox of checkboxes) {
                checkbox.checked = source.checked;
            }
        }

        function exportSelected() {
            const form = document.getElementById('exportForm');
            const checked = document.querySelectorAll('input[name="selected_item"]:checked');
            if (checked.length === 0) {
                alert("Por favor, selecciona al menos un activo para exportar.");
                return;
            }
            const ids = Array.from(checked).map(cb => cb.value).join(',');
            document.getElementById('selected_ids_input').value = ids;
            form.submit();
        }

        function exportAll() {
            const form = document.getElementById('exportForm');
            document.getElementById('selected_ids_input').value = ''; // Clear selection
            form.submit();
        }

        window.onclick = function(event) {
            if (event.target == modal) {
                closeModal();
            }
        }
    </script>
{% endblock %}
"""

# Configurar el cargador de plantillas de Jinja2 para usar las strings de HTML
app.jinja_loader = ChoiceLoader([
    DictLoader({
        'layout.html': HTML_LAYOUT,
        'login.html': HTML_LOGIN,
        'dashboard.html': HTML_DASHBOARD,
        'time_tracking.html': HTML_TIME_TRACKING,
        'forgot_clock.html': HTML_FORGOT_CLOCK,
        'admin_view.html': HTML_ADMIN_VIEW,
        'letter_generator.html': HTML_LETTER_GENERATOR,
        'inventory.html': HTML_INVENTORY
    }),
    app.jinja_loader,  # El cargador por defecto
])

# Configuración de Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = "Por favor, inicia sesión para acceder a esta página."
login_manager.login_message_category = "info"


# --- PASO 4: Base de Datos (SQLite) ---
def get_db():
    """Conecta a la base de datos."""
    db = sqlite3.connect(app.config['DATABASE'])
    db.row_factory = sqlite3.Row
    return db


def init_db():
    """Inicializa la base de datos y crea las tablas si no existen."""
    with app.app_context():
        db = get_db()
        cursor = db.cursor()
        cursor.execute('CREATE TABLE IF NOT EXISTS user (id INTEGER PRIMARY KEY, username TEXT UNIQUE NOT NULL)')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS time_entry (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                clock_in DATETIME NOT NULL,
                clock_out DATETIME,
                FOREIGN KEY (user_id) REFERENCES user (id)
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS balance_adjustment (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                adjustment_seconds INTEGER NOT NULL,
                reason TEXT NOT NULL,
                admin_user_id INTEGER NOT NULL,
                created_at DATETIME NOT NULL,
                FOREIGN KEY (user_id) REFERENCES user (id)
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS inventory (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                brand TEXT,
                color TEXT,
                quantity TEXT,
                status TEXT,
                location TEXT
            )
        ''')
        db.commit()
        print("Base de datos inicializada.")


def get_initial_inventory_data():
    """Retorna la lista completa de activos iniciales."""
    # Transcripción de datos de las imágenes proporcionadas
    return [
        {'name': 'Botas de Hule', 'brand': 'Varios', 'color': 'Negro', 'quantity': '7', 'status': 'Bueno',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Winch', 'brand': 'NP', 'color': 'Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Estante de Cochera'},
        {'name': 'Alfombra Azul', 'brand': '', 'color': 'Azul', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Pichinga de Gasolina Pequeña', 'brand': '', 'color': 'Rojo', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Pichinga de Gasolina Grande', 'brand': '', 'color': 'Naranja', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Aceite para Carro 10w30', 'brand': 'Repsol', 'color': 'Gris', 'quantity': '1',
         'status': 'Medio lleno', 'location': 'Estante Cochera Izquierda'},
        {'name': 'Cascos de proteccion', 'brand': 'NP', 'color': 'Azul, Blanco, Amarillo', 'quantity': '1',
         'status': '', 'location': 'Estante Cochera Izquierda'},
        {'name': 'Machete Kilimanjaro', 'brand': 'Kilimanjaro', 'color': 'Negro', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Hielera Pequeña', 'brand': '', 'color': 'Azul y Blanco', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Hielera Pequeña', 'brand': '', 'color': 'Azul y Naranja', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Chalecos Seguridad', 'brand': '', 'color': '', 'quantity': '', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Lingas ARB', 'brand': '', 'color': '', 'quantity': '', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Lingas CAT', 'brand': '', 'color': 'Amarillo', 'quantity': '', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Lingas sin marca', 'brand': '', 'color': '', 'quantity': '', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Zapatos de Seguridad', 'brand': '', 'color': 'Negro', 'quantity': '', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Pichingas Vacias', 'brand': '', 'color': '', 'quantity': '4', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Anillo Grande para PI', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Anillo Pequeño para PI', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Sonda', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Puntero', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Contador para Huincher', 'brand': '', 'color': 'Azul', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Silla Plegable', 'brand': 'ARB', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Izquierda'},
        {'name': 'Tripode Grande para Camara', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Barra de Tierra para RE', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Cables Amarillos para Tomo', 'brand': '', 'color': '', 'quantity': '22', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Lampara Amarilla', 'brand': '', 'color': 'Amarillo', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Abanicos con Luz', 'brand': '', 'color': 'Negro', 'quantity': '2', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Separador Cables en Cajuela', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Perra', 'brand': '', 'color': '', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Huincher', 'brand': '', 'color': 'Azul', 'quantity': '1', 'status': '',
         'location': 'Estante Cochera Derecha'},
        {'name': 'Generador Predator 3500', 'brand': '', 'color': '', 'quantity': '1', 'status': '', 'location': ''},
        {'name': 'Caja de Pines', 'brand': '', 'color': '', 'quantity': '1', 'status': '', 'location': ''},
        {'name': 'Mesa Plegable', 'brand': '', 'color': '', 'quantity': '1', 'status': '', 'location': ''},
        {'name': 'Hielera Grande', 'brand': '', 'color': 'Roja', 'quantity': '1', 'status': '',
         'location': 'Bodega Atrás'},
        {'name': 'Paraguas', 'brand': '', 'color': 'Negro con Blanco', 'quantity': '1', 'status': '', 'location': ''},
        {'name': 'Shield Guard', 'brand': 'Everlast', 'color': 'Camuflaje', 'quantity': '6', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Pines Metalicos', 'brand': 'NP', 'color': 'Acero Inoxidable', 'quantity': '110', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Machete', 'brand': 'Corneta', 'color': 'Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Mazo de Mano', 'brand': 'Surtek', 'color': 'Negro/Amarillo', 'quantity': '4', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Cinta Metrica 30m', 'brand': 'Stanley', 'color': 'Negro/Amarillo', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Cinta Metrica 50m', 'brand': 'Stanley', 'color': 'Amarillo', 'quantity': '2', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Cinta Metrica 50m', 'brand': 'Truper', 'color': 'Naranja/Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Cinta Metrica 100m', 'brand': 'Truper', 'color': 'Naranja/Negro', 'quantity': '2', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Cinta Metrica 100m', 'brand': 'Toolcraft', 'color': 'Gris', 'quantity': '2', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Cinta de Seguridad(Peligro)', 'brand': 'NP', 'color': 'Amarillo', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja de Pines'},
        {'name': 'Bolsa Multiusos', 'brand': 'Pricesmart', 'color': 'Amarillo y Verde', 'quantity': '2',
         'status': 'Bueno', 'location': 'Caja de Pines'},
        {'name': 'Triangulos de Seguridad', 'brand': 'Eurowarndreieck', 'color': 'Rojo', 'quantity': '2',
         'status': 'Bueno', 'location': 'Caja Ironman'},
        {'name': 'Lagartos Seguridad', 'brand': 'NP', 'color': 'Negro/Rojo', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Guantes de Cuero', 'brand': 'Truper', 'color': 'Naranja/Gris', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Caja de Herramientas', 'brand': 'Stack-On', 'color': 'Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Pala Plegable de Mano', 'brand': 'NP', 'color': 'Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Cable Adaptador de Energia para Carro', 'brand': 'NP', 'color': 'Rojo', 'quantity': '1',
         'status': 'Bueno', 'location': 'Caja Ironman'},
        {'name': 'Hacha de Mano', 'brand': 'Truper', 'color': 'Naranja/Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Mazo de Hule pequeño', 'brand': 'NP', 'color': 'Verde(despintado)', 'quantity': '1',
         'status': 'Bueno', 'location': 'Caja Ironman'},
        {'name': 'Cuchara para Cemento', 'brand': 'NP', 'color': 'Gris', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Tubos para Gata Mecanica', 'brand': 'NP', 'color': 'Plateado', 'quantity': '2', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Tubo Llave de Rana', 'brand': 'NP', 'color': 'Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Cubo de 17" y 19" (Reversible)', 'brand': 'NP', 'color': 'Plateado', 'quantity': '1',
         'status': 'Bueno', 'location': 'Caja Ironman'},
        {'name': 'Cubo de 21" y 23" (Reversible)', 'brand': 'NP', 'color': 'Plateado', 'quantity': '1',
         'status': 'Bueno', 'location': 'Caja Ironman'},
        {'name': 'Llave 19mm', 'brand': 'Eastman', 'color': 'Plateado', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Llave 17mm', 'brand': 'Eastman', 'color': 'Plateado', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Llave 14mm', 'brand': 'Eastman', 'color': 'Plateado', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Llave 13mm', 'brand': 'Eastman', 'color': 'Plateado', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Llave 11mm', 'brand': 'Eastman', 'color': 'Plateado', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Llave 9mm', 'brand': 'Eastman', 'color': 'Plateado', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Eslabon de Seguridad grande', 'brand': 'NP', 'color': 'Plateado', 'quantity': '4', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Eslabon de Seguridad pequeño', 'brand': 'NP', 'color': 'Plateado', 'quantity': '2', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Pines Sujetadores(anclaje) Toldo Carro', 'brand': 'NP', 'color': 'Plateado', 'quantity': '4',
         'status': 'Bueno', 'location': 'Caja Ironman'},
        {'name': 'Tensores de Cable', 'brand': 'NP', 'color': 'Plateado', 'quantity': '2', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Alicate', 'brand': 'Ineco', 'color': 'Negro/Amarillo', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Candado pequeño', 'brand': 'Yale', 'color': 'Dorado', 'quantity': '2', 'status': 'Buscar llaves',
         'location': 'Caja Ironman'},
        {'name': 'Candado pequeño', 'brand': 'CCP', 'color': 'Dorado', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Platina L', 'brand': 'NP', 'color': 'Negro', 'quantity': '1', 'status': 'Bueno',
         'location': 'Caja Ironman'},
        {'name': 'Base Monitor', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Bolso Bombilla Camara Pequeña', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Bolso Drone', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Bolso Extensiones', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Bombilla Camara Grande', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Caja Plastica Articulos Varios', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Caja Plastica Cables', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Camisas HAC', 'brand': '', 'color': '', 'quantity': '4', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Chaleco Geologia', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Divers', 'brand': 'Varios', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Grabador PYLE', 'brand': 'PYLE', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Kit Lodos', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Lector CD`s BluRay', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Revisar',
         'location': 'Closet Randall'},
        {'name': 'Mapas', 'brand': 'Varios', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Mecate Nylon', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Monitor de Grabador PYLE', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Niveles TopCon AT-G6', 'brand': '', 'color': '', 'quantity': '2', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Posters Ciencias Naturales', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Roseador con Sensor', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Sensor de Presion', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'SlimLogCom', 'brand': '', 'color': '', 'quantity': '2', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Sombreros HAC', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Tubo Anaranjado', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Zapatos CAT Punta Acero', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Zapatos Keen', 'brand': '', 'color': '', 'quantity': '1', 'status': 'Bueno',
         'location': 'Closet Randall'},
        {'name': 'Abrillantador de Focos', 'brand': 'RainX', 'color': 'Amarillo', 'quantity': '1',
         'status': 'Casi Lleno', 'location': 'Bodega del Medio'},
        {'name': 'Anti Fog', 'brand': 'RainX', 'color': 'Negro', 'quantity': '1', 'status': '1/2',
         'location': 'Bodega del Medio'},
        {'name': 'Cera Ceramica Liquida', 'brand': 'Meguiar`s', 'color': 'Azul', 'quantity': '1', 'status': '3/4',
         'location': 'Bodega del Medio'},
        {'name': 'Glass Cleaner & Rain Repellent', 'brand': 'RainX', 'color': 'Amarillo', 'quantity': '0',
         'status': '0', 'location': 'Bodega del Medio'},
        {'name': 'Limpiador de Asientos', 'brand': 'Weiman', 'color': 'Café', 'quantity': '2', 'status': 'Casi Llenos',
         'location': 'Bodega del Medio'},
        {'name': 'Limpiador de Vidrios', 'brand': 'RainX', 'color': 'Azul', 'quantity': '1', 'status': '1/4',
         'location': 'Bodega del Medio'},
        {'name': 'Limpiador y Abrillantador de Dash', 'brand': 'Meguiar`s', 'color': 'Negro', 'quantity': '1',
         'status': '1/16', 'location': 'Bodega del Medio'},
        {'name': 'Shampoo y Cera', 'brand': 'Meguiar`s', 'color': 'Amarillo', 'quantity': '1', 'status': '1/2',
         'location': 'Bodega del Medio'},
        {'name': 'Water Repellent Glass', 'brand': 'RainX', 'color': 'Amarillo', 'quantity': '1Gde 1Peq',
         'status': '1/2 y Lleno', 'location': 'Bodega del Medio'},
        {'name': 'Water Repellent Plastico', 'brand': 'RainX', 'color': 'Amarillo', 'quantity': '1', 'status': '1/4',
         'location': 'Bodega del Medio'},
        {'name': 'Wipes para Limpieza de Asientos', 'brand': 'Weiman', 'color': 'Café', 'quantity': '2',
         'status': 'Casi Llenos', 'location': 'Bodega del Medio'},
        {'name': 'Jabon Detergente Escobillas', 'brand': 'RainX', 'color': 'Amarillo', 'quantity': '1', 'status': '1/2',
         'location': 'Bodega del Medio'},
        {'name': 'Water Repellent Fast Max', 'brand': 'RainX', 'color': 'Azul', 'quantity': '1', 'status': '1/16',
         'location': 'Bodega del Medio'},
        {'name': 'Abrillantador de Llantas', 'brand': 'Meguiar`s', 'color': 'Negro', 'quantity': '1', 'status': 'Lleno',
         'location': 'Bodega del Medio'},
        {'name': 'Restaurador Partes Negras', 'brand': 'Meguiar`s', 'color': 'Negro', 'quantity': '1',
         'status': 'Lleno', 'location': 'Bodega del Medio'},
        {'name': 'Limpiador Dash extra Brillo', 'brand': 'Meguiar`s', 'color': 'Negro', 'quantity': '1',
         'status': 'Lleno', 'location': 'Bodega del Medio'},
        {'name': 'Agua Destilada Galon', 'brand': 'Quiflo', 'color': 'Azul', 'quantity': '2', 'status': 'Lleno',
         'location': 'Bodega del Medio'},
        {'name': 'Alcohol Gel Galon', 'brand': 'Xpro', 'color': 'Blanco', 'quantity': '1', 'status': '1/16',
         'location': 'Bodega del Medio'},
        {'name': 'Bolsas de Basura', 'brand': 'NP', 'color': 'Negro', 'quantity': '11', 'status': '3Gde 5Jdn 3Pqa',
         'location': 'Bodega del Medio'},
        {'name': 'Bolsas de Basura', 'brand': 'NP', 'color': 'Blanco', 'quantity': '1', 'status': 'Mitad',
         'location': 'Bodega del Medio'},
        {'name': 'Bombillos viejos casa', 'brand': 'NP', 'color': 'Blanco', 'quantity': '18', 'status': 'Usados',
         'location': 'Bodega del Medio'},
        {'name': 'Cepillo Grueso', 'brand': 'NP', 'color': 'Rosado', 'quantity': '1', 'status': 'Usado',
         'location': 'Bodega del Medio'},
        {'name': 'Cloro Galon', 'brand': 'Suli', 'color': 'Blanco', 'quantity': '1', 'status': '1/2',
         'location': 'Bodega del Medio'},
        {'name': 'Desinfectante Galon', 'brand': 'Bonito', 'color': 'Verde', 'quantity': '1', 'status': '1/2',
         'location': 'Bodega del Medio'},
        {'name': 'Desodorante Ambiental', 'brand': 'Sapolio', 'color': 'Varios', 'quantity': '7', 'status': 'Nuevo',
         'location': 'Bodega del Medio'},
        {'name': 'Escobita', 'brand': 'NP', 'color': 'Amarillo', 'quantity': '1', 'status': 'Nuevo',
         'location': 'Bodega del Medio'},
        {'name': 'Espuma Arranca Grasa', 'brand': 'Axion', 'color': 'Blanco', 'quantity': '1', 'status': 'Lleno',
         'location': 'Bodega del Medio'},
        {'name': 'Jabon en Polvo', 'brand': 'Xedex', 'color': 'Azul', 'quantity': '2,3 kg', 'status': '1KG',
         'location': 'Bodega del Medio'},
        {'name': 'Jabon Liquido Manos Galon', 'brand': 'Solimax', 'color': 'Naranja', 'quantity': '1', 'status': '1/4',
         'location': 'Bodega del Medio'},
        {'name': 'Lamparas LED', 'brand': 'EcoMax', 'color': 'Azul', 'quantity': '18', 'status': 'Nuevo',
         'location': 'Bodega del Medio'},
        {'name': 'Lava Platos Liquido Galon', 'brand': 'Briza', 'color': 'Verde', 'quantity': '1', 'status': 'Lleno',
         'location': 'Bodega del Medio'},
        {'name': 'Limpia Vidrios Galon', 'brand': 'MR Musculo', 'color': 'Blanco', 'quantity': '1',
         'status': 'Medio Lleno', 'location': 'Bodega del Medio'},
        {'name': 'Limpia Vidrios Spray', 'brand': 'Brasso', 'color': 'Blanco', 'quantity': '0', 'status': 'Vacio',
         'location': 'Bodega del Medio'},
        {'name': 'Limpia Vidrios Spray', 'brand': 'MR Musculo', 'color': 'Naranja/Azul', 'quantity': '1',
         'status': 'Medio Lleno', 'location': 'Bodega del Medio'},
        {'name': 'Limpiador de Superficies', 'brand': 'NP', 'color': 'Café', 'quantity': '1', 'status': 'Medio Lleno',
         'location': 'Bodega del Medio'},
        {'name': 'Limpiador de Sanitario', 'brand': 'Lysol', 'color': 'Azul', 'quantity': '1', 'status': 'Medio Vacio',
         'location': 'Bodega del Medio'},
        {'name': 'Lysol Spray', 'brand': 'Lysol', 'color': 'Azul', 'quantity': '1', 'status': 'Medio Lleno',
         'location': 'Bodega del Medio'},
        {'name': 'Papel Higienico Rollo Grande', 'brand': 'Jumbo Rollo', 'color': 'Morado', 'quantity': '4',
         'status': 'Nuevo', 'location': 'Bodega del Medio'},
        {'name': 'Potasa', 'brand': 'Potasa', 'color': 'Blanco', 'quantity': '0', 'status': 'Nuevo',
         'location': 'Bodega del Medio'},
        {'name': 'Removedor de Grasa', 'brand': 'EcoLife', 'color': 'Blanco', 'quantity': '1', 'status': '1/2',
         'location': 'Bodega del Medio'},
        {'name': 'Removedor de Grasa', 'brand': 'Ultra Cleaner', 'color': 'Blanco', 'quantity': '0', 'status': 'Vacio',
         'location': 'Bodega del Medio'},
        {'name': 'Sanitizante Galon', 'brand': 'Adarga', 'color': 'Blanco', 'quantity': '1', 'status': '1/4',
         'location': 'Bodega del Medio'},
        {'name': 'Servilleta limpiar manos', 'brand': 'Suavecell', 'color': 'Azul', 'quantity': '5', 'status': 'Nuevo',
         'location': 'Bodega del Medio'},
        {'name': 'Vinagre Blanco Galon', 'brand': 'La Unica', 'color': 'Blanco', 'quantity': '1', 'status': '1/5',
         'location': 'Bodega del Medio'},
    ]


def populate_initial_inventory():
    """Popula la tabla de inventario con los datos iniciales si está vacía."""
    with app.app_context():
        db = get_db()
        cursor = db.cursor()
        cursor.execute("SELECT COUNT(id) FROM inventory")
        count = cursor.fetchone()[0]
        if count == 0:
            print(">>> Populando inventario inicial...")
            initial_data = get_initial_inventory_data()
            cursor.executemany(
                "INSERT INTO inventory (name, brand, color, quantity, status, location) VALUES (:name, :brand, :color, :quantity, :status, :location)",
                initial_data
            )
            db.commit()
            print(f">>> {len(initial_data)} activos agregados al inventario.")


# --- PASO 5: Modelo de Usuario y Datos ---
class User(UserMixin):
    def __init__(self, id, username, password_hash, daily_hours):
        self.id = id
        self.username = username
        self.password = password_hash
        self.daily_hours = daily_hours


USERS = {
    '1': {'username': 'Admin', 'password': generate_password_hash('Rindy'), 'daily_hours': 0},
    '2': {'username': 'Esteban', 'password': generate_password_hash('HAC127@'), 'daily_hours': 9.6},
    '3': {'username': 'Hernan', 'password': generate_password_hash('Obi123'), 'daily_hours': 9.6},
    '4': {'username': 'Jervaice', 'password': generate_password_hash('Tortas'), 'daily_hours': 8.0}
}


@login_manager.user_loader
def load_user(user_id):
    user_data = USERS.get(str(user_id))
    if user_data:
        return User(user_id, user_data['username'], user_data['password'], user_data['daily_hours'])
    return None


# --- PASO 6: Lógica de Negocio ---
def get_user_by_username(username):
    for user_id, user_data in USERS.items():
        if user_data['username'] == username:
            return User(user_id, user_data['username'], user_data['password'], user_data['daily_hours'])
    return None


def seconds_to_str(seconds):
    if seconds is None: return ""
    sign = "-" if seconds < 0 else "+"
    seconds = abs(seconds)
    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{sign}{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}"


def parse_balance_string(balance_str):
    balance_str = balance_str.strip()
    match = re.match(r'([+-])?\s*(\d+):(\d+)', balance_str)
    if not match:
        raise ValueError("Formato de balance inválido. Use +HH:MM o -HH:MM.")
    sign, hours, minutes = match.groups()
    total_seconds = int(hours) * 3600 + int(minutes) * 60
    return -total_seconds if sign == '-' else total_seconds


def calculate_balance(user_id, db_cursor):
    user = load_user(str(user_id))
    if not user or user.daily_hours == 0:
        return "N/A", 0
    db_cursor.execute("SELECT clock_in, clock_out FROM time_entry WHERE user_id = ? AND clock_out IS NOT NULL",
                      (user_id,))
    entries = db_cursor.fetchall()
    total_worked_seconds = sum(
        (datetime.fromisoformat(e['clock_out']) - datetime.fromisoformat(e['clock_in'])).total_seconds() for e in
        entries)
    work_days = {datetime.fromisoformat(e['clock_in']).date() for e in entries if
                 datetime.fromisoformat(e['clock_in']).weekday() < 5}
    required_seconds = len(work_days) * user.daily_hours * 3600
    db_cursor.execute("SELECT SUM(adjustment_seconds) FROM balance_adjustment WHERE user_id = ?", (user_id,))
    total_adjustment_seconds = db_cursor.fetchone()[0] or 0
    balance_seconds = (total_worked_seconds - required_seconds) + total_adjustment_seconds
    hours, remainder = divmod(abs(balance_seconds), 3600)
    minutes, _ = divmod(remainder, 60)
    sign = "" if balance_seconds >= 0 else "-"
    balance_str = f"{sign}{int(hours):02d}h {int(minutes):02d}m"
    return balance_str, balance_seconds


# --- PASO 8: Rutas de la Aplicación (Controladores) ---

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('login.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = get_user_by_username(username)
        if user and check_password_hash(user.password, password):
            login_user(user)
            flash(f'Bienvenido de nuevo, {user.username}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Usuario o contraseña incorrectos.', 'danger')
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Has cerrado sesión correctamente.', 'info')
    return redirect(url_for('login'))


@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')


@app.route('/time_tracking')
@login_required
def time_tracking():
    if current_user.username == 'Admin':
        return redirect(url_for('admin_view'))
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT id, clock_in, clock_out FROM time_entry WHERE user_id = ? ORDER BY clock_in DESC",
                   (current_user.id,))
    raw_entries = cursor.fetchall()
    entries = []
    for row in raw_entries:
        entry = dict(row)
        entry['clock_in'] = datetime.fromisoformat(entry['clock_in'])
        if entry['clock_out']:
            entry['clock_out'] = datetime.fromisoformat(entry['clock_out'])
        entries.append(entry)
    active_entry = next((e for e in entries if e['clock_out'] is None), None)
    balance_str, balance_seconds = calculate_balance(current_user.id, cursor)
    return render_template('time_tracking.html', entries=entries, active_entry=active_entry, balance_str=balance_str,
                           balance_seconds=balance_seconds)


@app.route('/clock_in', methods=['POST'])
@login_required
def clock_in():
    db = get_db()
    cursor = db.cursor()
    cursor.execute("INSERT INTO time_entry (user_id, clock_in) VALUES (?, ?)", (current_user.id, datetime.now()))
    db.commit()
    flash('Entrada marcada correctamente.', 'success')
    return redirect(url_for('time_tracking'))


@app.route('/clock_out', methods=['POST'])
@login_required
def clock_out():
    db = get_db()
    cursor = db.cursor()
    cursor.execute("UPDATE time_entry SET clock_out = ? WHERE user_id = ? AND clock_out IS NULL",
                   (datetime.now(), current_user.id))
    db.commit()
    flash('Salida marcada correctamente.', 'success')
    return redirect(url_for('time_tracking'))


@app.route('/forgot_clock', methods=['GET'])
@login_required
def forgot_clock():
    if current_user.username not in ['Esteban', 'Hernan']:
        flash('No tienes permiso para acceder a esta función.', 'danger')
        return redirect(url_for('time_tracking'))
    return render_template('forgot_clock.html')


@app.route('/save_forgotten_clock', methods=['POST'])
@login_required
def save_forgotten_clock():
    if current_user.username not in ['Esteban', 'Hernan']:
        return redirect(url_for('time_tracking'))
    entry_date = request.form['date']
    clock_in_time = request.form['clock_in_time']
    clock_out_time = request.form['clock_out_time']
    try:
        clock_in_dt = datetime.fromisoformat(f"{entry_date}T{clock_in_time}")
        clock_out_dt = datetime.fromisoformat(f"{entry_date}T{clock_out_time}")
        if clock_out_dt <= clock_in_dt:
            flash('La hora de salida debe ser posterior a la de entrada.', 'danger')
            return redirect(url_for('forgot_clock'))
        db = get_db()
        cursor = db.cursor()
        cursor.execute("INSERT INTO time_entry (user_id, clock_in, clock_out) VALUES (?, ?, ?)",
                       (current_user.id, clock_in_dt, clock_out_dt))
        db.commit()
        flash('Marcaje manual añadido correctamente.', 'success')
    except ValueError:
        flash('Formato de fecha u hora inválido.', 'danger')
        return redirect(url_for('forgot_clock'))
    return redirect(url_for('time_tracking'))


@app.route('/admin')
@login_required
def admin_view():
    if current_user.username != 'Admin':
        flash('Acceso no autorizado.', 'danger')
        return redirect(url_for('dashboard'))
    users_info = {}
    db = get_db()
    cursor = db.cursor()
    for user_id_str, user_data in USERS.items():
        if user_data['username'] == 'Admin': continue
        user_id = int(user_id_str)
        balance_str, balance_seconds = calculate_balance(user_id, cursor)
        cursor.execute("SELECT * FROM time_entry WHERE user_id = ? ORDER BY clock_in DESC", (user_id,))
        raw_entries = cursor.fetchall()
        entries = [dict(row) for row in raw_entries]
        for e in entries:
            e['clock_in'] = datetime.fromisoformat(e['clock_in'])
            if e['clock_out']: e['clock_out'] = datetime.fromisoformat(e['clock_out'])
        cursor.execute("SELECT * FROM balance_adjustment WHERE user_id = ? ORDER BY created_at DESC", (user_id,))
        raw_adjustments = cursor.fetchall()
        adjustments = []
        for row in raw_adjustments:
            adj = dict(row)
            adj['created_at'] = datetime.fromisoformat(adj['created_at'])
            adj['adjustment_str'] = seconds_to_str(adj['adjustment_seconds'])
            admin_user = load_user(adj['admin_user_id'])
            adj['admin_username'] = admin_user.username if admin_user else "Desconocido"
            adjustments.append(adj)
        users_info[user_id] = {'username': user_data['username'], 'balance_str': balance_str,
                               'balance_seconds': balance_seconds, 'entries': entries, 'adjustments': adjustments}
    return render_template('admin_view.html', users_info=users_info)


@app.route('/admin/edit/<int:entry_id>', methods=['POST'])
@login_required
def admin_edit_entry(entry_id):
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    clock_in = request.form.get('clock_in')
    clock_out = request.form.get('clock_out')
    db = get_db()
    cursor = db.cursor()
    cursor.execute("UPDATE time_entry SET clock_in = ?, clock_out = ? WHERE id = ?",
                   (clock_in, clock_out if clock_out else None, entry_id))
    db.commit()
    flash('Marcaje actualizado.', 'success')
    return redirect(url_for('admin_view'))


@app.route('/admin/delete/<int:entry_id>')
@login_required
def admin_delete_entry(entry_id):
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    db = get_db()
    cursor = db.cursor()
    cursor.execute("DELETE FROM time_entry WHERE id = ?", (entry_id,))
    db.commit()
    flash('Marcaje eliminado.', 'success')
    return redirect(url_for('admin_view'))


@app.route('/admin/set_balance/<int:user_id>', methods=['POST'])
@login_required
def admin_set_balance(user_id):
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    new_total_balance_str = request.form.get('new_total_balance')
    reason = request.form.get('reason')
    if not reason:
        flash('Se requiere un motivo para el ajuste de balance.', 'danger')
        return redirect(url_for('admin_view'))
    try:
        new_total_seconds = parse_balance_string(new_total_balance_str)
    except ValueError as e:
        flash(str(e), 'danger')
        return redirect(url_for('admin_view'))
    db = get_db()
    cursor = db.cursor()
    _, current_balance_seconds = calculate_balance(user_id, cursor)
    adjustment_needed = new_total_seconds - current_balance_seconds
    cursor.execute(
        "INSERT INTO balance_adjustment (user_id, adjustment_seconds, reason, admin_user_id, created_at) VALUES (?, ?, ?, ?, ?)",
        (user_id, adjustment_needed, reason, current_user.id, datetime.now()))
    db.commit()
    flash(f"Balance para el usuario actualizado correctamente.", 'success')
    return redirect(url_for('admin_view'))


@app.route('/admin/letter_generator')
@login_required
def letter_generator():
    if current_user.username != 'Admin':
        flash('Acceso no autorizado.', 'danger')
        return redirect(url_for('dashboard'))
    today_date = date.today().strftime('%Y-%m-%d')
    current_year = date.today().year
    return render_template('letter_generator.html', today_date=today_date, current_year=current_year)


@app.route('/admin/generate_doc', methods=['POST'])
@login_required
def generate_doc():
    if current_user.username != 'Admin': return redirect(url_for('dashboard'))
    data = {k: request.form.get(k) for k in
            ['doc_date', 'project_name', 'client_name', 'year', 'contact_person', 'contact_position', 'contact_email',
             'project_type']}
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    try:
        date_obj = datetime.strptime(data['doc_date'], '%Y-%m-%d')
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre",
                 "noviembre", "diciembre"]
        formatted_date = f"{date_obj.day} de {meses[date_obj.month - 1]} de {date_obj.year}"
    except (ValueError, TypeError):
        formatted_date = data['doc_date']
    p_date = document.add_paragraph(f'Fecha: {formatted_date}')
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()
    document.add_paragraph('A quien interese:', style='Normal').add_run().bold = True
    document.add_paragraph(
        'Por medio de la presente, queremos expresar nuestra satisfacción con los trabajos realizados por HidroAmbiente Consultores.')
    document.add_paragraph(
        'Su profesionalismo y compromiso han sido fundamentales para el éxito de los proyectos desarrollados, cumpliendo con nuestras expectativas y necesidades.')
    document.add_paragraph('A continuación, se presenta un cuadro detallado con la información correspondiente:')
    document.add_paragraph()
    table_data = [("Tipo de Proyecto", data['project_type']), ("Nombre del Proyecto", data['project_name']),
                  ("Nombre del Cliente", data['client_name']), ("Año", data['year']),
                  ("Contacto", data['contact_person']), ("Puesto del Contacto", data['contact_position']),
                  ("Correo Electrónico", data['contact_email'])]
    table = document.add_table(rows=len(table_data), cols=2)
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(table_data):
        table.cell(i, 0).text = key
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value
    document.add_paragraph()
    document.add_paragraph(
        'Finalmente, agradecemos nuevamente su excelente trabajo y quedamos en espera de futuras colaboraciones que sigan fortaleciendo esta relación profesional.')
    document.add_paragraph()
    document.add_paragraph('Atentamente,')
    document.add_paragraph('\n\n')
    document.add_paragraph('_________________________')
    document.add_paragraph(data['contact_person']).add_run().bold = True
    document.add_paragraph(data['contact_position'])
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name=f"Carta_Satisfaccion_{data['project_name']}.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# --- Rutas para Inventario ---
@app.route('/inventory')
@login_required
def inventory_view():
    if current_user.username not in ['Admin', 'Jervaice']:
        flash('Acceso no autorizado.', 'danger')
        return redirect(url_for('dashboard'))
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT * FROM inventory ORDER BY name")
    items = cursor.fetchall()
    return render_template('inventory.html', items=items)


@app.route('/inventory/add', methods=['POST'])
@login_required
def add_inventory_item():
    if current_user.username not in ['Admin', 'Jervaice']: return jsonify({'error': 'Unauthorized'}), 403
    db = get_db()
    cursor = db.cursor()
    cursor.execute(
        "INSERT INTO inventory (name, brand, color, quantity, status, location) VALUES (?, ?, ?, ?, ?, ?)",
        (request.form.get('name'), request.form.get('brand'), request.form.get('color'), request.form.get('quantity'),
         request.form.get('status'), request.form.get('location'))
    )
    db.commit()
    flash('Activo añadido correctamente.', 'success')
    return redirect(url_for('inventory_view'))


@app.route('/inventory/edit/<int:item_id>', methods=['POST'])
@login_required
def edit_inventory_item(item_id):
    if current_user.username not in ['Admin', 'Jervaice']: return jsonify({'error': 'Unauthorized'}), 403
    db = get_db()
    cursor = db.cursor()
    cursor.execute(
        "UPDATE inventory SET name=?, brand=?, color=?, quantity=?, status=?, location=? WHERE id=?",
        (request.form.get('name'), request.form.get('brand'), request.form.get('color'), request.form.get('quantity'),
         request.form.get('status'), request.form.get('location'), item_id)
    )
    db.commit()
    flash('Activo actualizado correctamente.', 'success')
    return redirect(url_for('inventory_view'))


@app.route('/inventory/delete/<int:item_id>')
@login_required
def delete_inventory_item(item_id):
    if current_user.username not in ['Admin', 'Jervaice']: return redirect(url_for('dashboard'))
    db = get_db()
    cursor = db.cursor()
    cursor.execute("DELETE FROM inventory WHERE id = ?", (item_id,))
    db.commit()
    flash('Activo eliminado correctamente.', 'success')
    return redirect(url_for('inventory_view'))


@app.route('/inventory/export', methods=['POST'])
@login_required
def export_inventory():
    if current_user.username not in ['Admin', 'Jervaice']: return redirect(url_for('dashboard'))

    export_type = request.form.get('export_type')
    db = get_db()
    cursor = db.cursor()

    if export_type == 'selected':
        selected_ids_str = request.form.get('selected_ids', '').split(',')
        if not selected_ids_str or selected_ids_str == ['']:
            flash('No se seleccionaron activos para exportar.', 'warning')
            return redirect(url_for('inventory_view'))

        selected_ids = [int(i) for i in selected_ids_str if i]
        placeholders = ','.join('?' for i in selected_ids)
        query = f"SELECT * FROM inventory WHERE id IN ({placeholders}) ORDER BY name"
        cursor.execute(query, selected_ids)
    else:  # export all
        cursor.execute("SELECT * FROM inventory ORDER BY name")

    items = cursor.fetchall()

    # Create Excel file in memory
    wb = Workbook()
    ws = wb.active
    ws.title = "Inventario de Activos"

    # Header
    headers = ["Activo", "Marca", "Color", "Cantidad", "Estado", "Ubicación"]
    ws.append(headers)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')

    # Data
    for item in items:
        ws.append([item['name'], item['brand'], item['color'], item['quantity'], item['status'], item['location']])

    # Save to a stream
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"Inventario_HAC2025_{date.today().isoformat()}.xlsx",
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


# --- PASO 9: Punto de Entrada Principal ---
if __name__ == '__main__':
    if not os.path.exists(app.config['DATABASE']):
        print(f"No se encontró la base de datos '{app.config['DATABASE']}'. Creando...")
        init_db()

    # Populate inventory only if it's empty
    populate_initial_inventory()

    print("\n>>> Iniciando servidor de desarrollo de Flask...")
    print(f">>> Abre tu navegador y ve a http://127.0.0.1:5000")
    app.run(debug=True)
